from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement

from indexer.extractor import extract


def test_docx_explicit_page_breaks_set_chunk_pages(tmp_path):
    path = tmp_path / "paged.docx"
    doc = Document()
    doc.add_paragraph("alpha first page")
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    paragraph.add_run("beta second page")
    doc.save(path)

    chunks = extract(path)

    assert [(chunk["page"], chunk["text"]) for chunk in chunks] == [
        (1, "alpha first page"),
        (2, "beta second page"),
    ]


def test_docx_rendered_page_breaks_set_chunk_pages(tmp_path):
    path = tmp_path / "rendered-break.docx"
    doc = Document()
    doc.add_paragraph("alpha first page")
    run = doc.add_paragraph().add_run()
    run._r.append(OxmlElement("w:lastRenderedPageBreak"))
    run.add_text("beta rendered second page")
    doc.save(path)

    chunks = extract(path)

    assert [(chunk["page"], chunk["text"]) for chunk in chunks] == [
        (1, "alpha first page"),
        (2, "beta rendered second page"),
    ]
