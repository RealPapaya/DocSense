from docx import Document
from docx.enum.text import WD_BREAK
from openpyxl import Workbook

from indexer.extractor import extract


def test_docx_explicit_page_breaks_set_chunk_pages(tmp_path):
    path = tmp_path / "paged.docx"
    doc = Document()
    doc.add_paragraph("alpha first page")
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    paragraph.add_run("beta second page")
    doc.save(path)

    chunks = extract(path)

    assert [(chunk["page"], chunk["text"]) for chunk in chunks] == [
        (1, "alpha first page"),
        (2, "beta second page"),
    ]


def test_xlsx_uses_sheet_name_as_page_label(tmp_path):
    path = tmp_path / "sheets.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "SpecTable"
    ws.append(["BIOS", "UEFI"])
    ws2 = wb.create_sheet("Config")
    ws2.append(["Power", "Enabled"])
    wb.save(path)

    chunks = extract(path)

    assert [(chunk["page"], chunk["text"]) for chunk in chunks] == [
        ("SpecTable", "[Sheet: SpecTable]\nBIOS\tUEFI"),
        ("Config", "[Sheet: Config]\nPower\tEnabled"),
    ]
